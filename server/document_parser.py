import sys
import os
import json
import re
import io

# Force UTF-8 output encoding for Windows shells and child_process
if sys.stdout.encoding != 'utf-8':
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')


def parse_docx(file_path):
    import docx
    doc = docx.Document(file_path)
    
    html_parts = []
    has_table = len(doc.tables) > 0
    table_cols = []
    
    # Process elements in order
    title = ""
    
    # Extract paragraphs
    for p in doc.paragraphs:
        text = p.text.strip()
        if not text:
            continue
        
        if not title and any(k in text.upper() for k in ["ANNEXURE", "SCHEDULE", "DECLARATION", "UNDERTAKING", "FORM", "PROFORMA", "CERTIFICATE"]):
            title = text
            
        # Check if heading
        is_heading = p.style and ('Heading' in p.style.name or 'Title' in p.style.name)
        is_bold_title = False
        if len(p.runs) > 0 and all(r.bold for r in p.runs if r.text.strip()):
            is_bold_title = True
            
        if is_heading or is_bold_title or text.upper().startswith("ANNEXURE") or "DECLARATION" in text.upper():
            html_parts.append(f'<h3 style="font-size: 15px; font-weight: bold; text-align: center; text-transform: uppercase; margin: 14px 0 8px 0; text-decoration: underline;">{text}</h3>')
        else:
            # Inline formatting
            p_html = ""
            for r in p.runs:
                r_text = r.text
                if not r_text: continue
                if r.bold:
                    r_text = f"<strong>{r_text}</strong>"
                if r.italic:
                    r_text = f"<em>{r_text}</em>"
                if r.underline:
                    r_text = f"<u>{r_text}</u>"
                p_html += r_text
            
            align = "left"
            if p.alignment == docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER:
                align = "center"
            elif p.alignment == docx.enum.text.WD_ALIGN_PARAGRAPH.RIGHT:
                align = "right"
            elif p.alignment == docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY:
                align = "justify"
                
            html_parts.append(f'<p style="margin-bottom: 8px; line-height: 1.5; font-size: 11pt; text-align: {align};">{p_html or text}</p>')
            
    # Extract tables
    for t_idx, tbl in enumerate(doc.tables):
        tbl_html = ['<table style="width: 100%; border-collapse: collapse; margin: 14px 0; font-size: 10pt;">']
        
        # Header row
        if len(tbl.rows) > 0:
            header_cells = [c.text.strip() for c in tbl.rows[0].cells]
            if t_idx == 0:
                table_cols = [{"key": f"col_{i}", "label": col} for i, col in enumerate(header_cells)]
            tbl_html.append('<thead><tr>')
            for cell in tbl.rows[0].cells:
                tbl_html.append(f'<th style="border: 1px solid black; padding: 6px 8px; background-color: #f0f0f0; font-weight: bold; text-align: center;">{cell.text.strip()}</th>')
            tbl_html.append('</tr></thead><tbody>')
            
            for row in tbl.rows[1:]:
                tbl_html.append('<tr>')
                for cell in row.cells:
                    tbl_html.append(f'<td style="border: 1px solid black; padding: 6px 8px;">{cell.text.strip()}</td>')
                tbl_html.append('</tr>')
            tbl_html.append('</tbody></table>')
            html_parts.append(''.join(tbl_html))
            
    if not title:
        title = os.path.splitext(os.path.basename(file_path))[0].replace("_", " ").title()
        
    return {
        "title": title[:100],
        "contentTemplate": "\n".join(html_parts),
        "isTable": has_table,
        "tableColumns": table_cols
    }

def parse_pdf(file_path):
    import pypdf
    reader = pypdf.PdfReader(file_path)
    pages_text = []
    for page in reader.pages:
        txt = page.extract_text()
        if txt:
            pages_text.append(txt)
            
    full_text = "\n\n".join(pages_text)
    paragraphs = full_text.split("\n\n")
    
    html_parts = []
    title = ""
    
    for p_text in paragraphs:
        cleaned = p_text.strip()
        if not cleaned: continue
        
        lines = cleaned.split("\n")
        first_line = lines[0].strip()
        
        if not title and any(k in first_line.upper() for k in ["ANNEXURE", "SCHEDULE", "DECLARATION", "UNDERTAKING", "FORM", "PROFORMA", "CERTIFICATE"]):
            title = first_line
            
        if any(k in first_line.upper() for k in ["ANNEXURE", "SCHEDULE", "DECLARATION", "UNDERTAKING", "FORM", "PROFORMA", "CERTIFICATE"]) and len(first_line) < 100:
            html_parts.append(f'<h3 style="font-size: 15px; font-weight: bold; text-align: center; text-transform: uppercase; margin: 14px 0 8px 0; text-decoration: underline;">{first_line}</h3>')
            remaining = "\n".join(lines[1:]).strip()
            if remaining:
                html_parts.append(f'<p style="margin-bottom: 8px; line-height: 1.5; font-size: 11pt; text-align: justify;">{remaining.replace(chr(10), " ")}</p>')
        else:
            html_parts.append(f'<p style="margin-bottom: 8px; line-height: 1.5; font-size: 11pt; text-align: justify;">{cleaned.replace(chr(10), " ")}</p>')
            
    if not title:
        title = os.path.splitext(os.path.basename(file_path))[0].replace("_", " ").title()
        
    return {
        "title": title[:100],
        "contentTemplate": "\n".join(html_parts),
        "isTable": False,
        "tableColumns": []
    }

def parse_txt(file_path):
    with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
        text = f.read()
    
    paragraphs = [p.strip() for p in text.split('\n\n') if p.strip()]
    html_parts = []
    title = ""
    
    for p in paragraphs:
        if not title:
            title = p.split('\n')[0][:80]
        html_parts.append(f'<p style="margin-bottom: 8px; line-height: 1.5; font-size: 11pt; text-align: justify;">{p.replace(chr(10), "<br/>")}</p>')
        
    return {
        "title": title or "Custom Annexure",
        "contentTemplate": "\n".join(html_parts),
        "isTable": False,
        "tableColumns": []
    }

def main():
    if len(sys.argv) < 2:
        print(json.dumps({"error": "No file path provided"}))
        sys.exit(1)
        
    file_path = sys.argv[1]
    if not os.path.exists(file_path):
        print(json.dumps({"error": f"File not found: {file_path}"}))
        sys.exit(1)
        
    ext = os.path.splitext(file_path)[1].lower()
    
    try:
        if ext in ['.docx', '.doc']:
            result = parse_docx(file_path)
        elif ext == '.pdf':
            result = parse_pdf(file_path)
        elif ext in ['.txt', '.md']:
            result = parse_txt(file_path)
        else:
            # Fallback text
            result = parse_txt(file_path)
            
        print(json.dumps(result, ensure_ascii=False))
    except Exception as e:
        print(json.dumps({"error": str(e)}))
        sys.exit(1)

if __name__ == '__main__':
    main()
