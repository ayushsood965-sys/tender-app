import docx
import zipfile
import sys
import re

sys.stdout.reconfigure(encoding='utf-8')

def inspect_sports_tender(docx_path):
    print("=" * 80)
    print(f"ANALYZING: {docx_path}")
    print("=" * 80)
    
    doc = docx.Document(docx_path)
    
    print(f"Total Paragraphs: {len(doc.paragraphs)}")
    print(f"Total Tables: {len(doc.tables)}")
    print(f"Total Sections: {len(doc.sections)}")
    
    print("\n--- ALL PARAGRAPHS WITH STYLES ---")
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if text:
            print(f"P[{i}] ({p.style.name}) [align={p.alignment}]: {text}")
            
    print("\n--- ALL TABLES ---")
    for t_idx, table in enumerate(doc.tables):
        print(f"\n--- TABLE {t_idx + 1} ({len(table.rows)} rows x {len(table.columns)} cols) ---")
        for r_idx, row in enumerate(table.rows):
            cells_text = [c.text.strip().replace('\n', ' | ') for c in row.cells]
            print(f"  Row {r_idx + 1}: {' /// '.join(cells_text)}")

inspect_sports_tender(r"C:\Users\Ayush Sood\Downloads\sports final tender.docx")
