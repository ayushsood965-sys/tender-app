import requests
import json
import os
import docx
import sys
import io

if sys.stdout.encoding != 'utf-8':
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

BASE_URL = "http://localhost:5000/api"

def test_full_annexure_flow():
    print("==================================================")
    print("1. TESTING DASHBOARD & ANNEXURES API")
    print("==================================================")
    r = requests.get(f"{BASE_URL}/dashboard-data")
    assert r.status_code == 200, f"Dashboard data failed: {r.text}"
    data = r.json()
    
    categories = data.get("categories", [])
    terms = data.get("terms", [])
    annexures = data.get("annexures", [])
    
    print(f"✅ Loaded {len(categories)} categories, {len(terms)} terms, {len(annexures)} annexures")
    assert len(annexures) >= 15, "Expected at least 15 master annexures"
    
    # Check term linkages
    mapped_terms = [t for t in terms if t.get("hasAnnexure")]
    print(f"✅ Found {len(mapped_terms)} terms linked to Annexures")
    for mt in mapped_terms[:5]:
        print(f"   ↳ Term '{mt['title']}' -> Annexure ID {mt['annexureId']}: '{mt.get('annexureTitle')}'")
        
    print("\n==================================================")
    print("2. TESTING TENDER CREATION WITH MANDATORY ANNEXURES")
    print("==================================================")
    
    # A. Create Limited Tender
    lt_payload = {
        "documentType": "Limited Tender Document",
        "tenderType": "limited",
        "tenderName": "Purchase of Sports Equipment for Hostels",
        "tenderNo": "1-13/2009-HPU (CW)",
        "departmentName": "Store Purchase Office",
        "departmentEmail": "spo@hpuniv.ac.in",
        "tenderInvitingAuthority": "Store Purchase Officer",
        "procurementLocations": "Tagore Boys Hostel, SBS Tribal Boys Hostel, and Dr. Y.S. Parmar Boys Hostel",
        "estimatedCost": 229000,
        "estimatedCostWords": "Rupees Two Lakh Twenty-Nine Thousand only",
        "emdAmount": 5000,
        "emdAmountWords": "Rupees Five Thousand Only",
        "pbgAmount": 10000,
        "pbgAmountWords": "Rupees Ten Thousand Only",
        "deliveryDays": 15,
        "courtJurisdiction": "Courts in Shimla Urban",
        "isAnnexureRequired": True,
        "selectedAnnexureIds": [1, 14], # Annexure-A Financial Bid + Non-Blacklisting
        "itemsList": [
            { "srNo": 1, "name": "Table Tennis Table", "specs": "22mm Top, TTFI approved", "quantity": "3" },
            { "srNo": 2, "name": "Badminton Racket", "specs": "Carbon material, Lightweight", "quantity": "10" }
        ],
        "selectedTermIds": [102, 408, 1303]
    }
    
    # Python boolean fix
    lt_payload["isAnnexureRequired"] = True
    
    r_lt = requests.post(f"{BASE_URL}/tenders", json=lt_payload)
    assert r_lt.status_code == 200, f"Failed to create Limited Tender: {r_lt.text}"
    created_lt = r_lt.json()
    print(f"✅ Created Limited Tender ID: {created_lt['id']}")
    
    # B. Create GeM ATC Tender
    gem_payload = {
        "documentType": "GeM ATC Document",
        "tenderType": "gem",
        "tenderName": "Procurement of High Resolution XRD System",
        "tenderNo": "GEM/2026/B/9876543",
        "departmentName": "Department of Physics",
        "departmentEmail": "physicshpu@gmail.com",
        "tenderInvitingAuthority": "Chairman, Department of Physics",
        "itemQuantity": "1 Complete System",
        "estimatedCost": 15000000,
        "estimatedCostWords": "Rupees One Crore Fifty Lakhs only",
        "emdAmount": 300000,
        "emdAmountWords": "Rupees Three Lakhs Only",
        "pbgAmount": 750000,
        "pbgAmountWords": "Rupees Seven Lakh Fifty Thousand Only",
        "isAnnexureRequired": True,
        "selectedAnnexureIds": [2, 3, 4, 5, 6, 7, 8, 9, 10, 11],
        "selectedTermIds": [102, 104, 302, 306, 401, 402, 408, 701, 1303, 1901]
    }
    r_gem = requests.post(f"{BASE_URL}/tenders", json=gem_payload)
    assert r_gem.status_code == 200, f"Failed to create GeM ATC: {r_gem.text}"
    created_gem = r_gem.json()
    print(f"✅ Created GeM ATC Tender ID: {created_gem['id']}")
    
    print("\n==================================================")
    print("3. TESTING DOCUMENT SAVING & DOCX EXPORT")
    print("==================================================")
    
    # Simulate saving document versions with annexure pages
    save_payload = {
        "tenderId": created_lt["id"],
        "name": "Limited Tender Final Submission",
        "pages": {
            "1": "<h2>Himachal Pradesh University</h2><p>Limited Tender Notice No. 1-13/2009-HPU (CW)</p>",
            "2": "<h3>Instructions to Bidders</h3><p>EMD: Rs. 5,000/-</p>",
            "3": "<h2>Annexure-A</h2><h3>Financial Bid Submission Form</h3><table><tr><th>Sr. No.</th><th>Name of Articles</th><th>Specification</th><th>Total Qty</th><th>Base Rate</th><th>GST (%)</th><th>Total Amount</th></tr><tr><td>1.</td><td>Table Tennis Table</td><td>22mm Top</td><td>3</td><td></td><td></td><td></td></tr></table>"
        }
    }
    r_save = requests.post(f"{BASE_URL}/documents", json=save_payload)
    assert r_save.status_code == 200, f"Failed to save document: {r_save.text}"
    saved_doc = r_save.json()
    print(f"✅ Saved Document ID: {saved_doc['id']}")
    
    # Test Word DOCX download
    r_docx = requests.get(f"{BASE_URL}/documents/{saved_doc['id']}/docx")
    assert r_docx.status_code == 200, f"Failed to generate DOCX: {r_docx.status_code}"
    assert len(r_docx.content) > 1000, "DOCX file size suspiciously small"
    
    temp_docx_path = "scratch_test_export.docx"
    with open(temp_docx_path, "wb") as f:
        f.write(r_docx.content)
        
    doc = docx.Document(temp_docx_path)
    print(f"✅ Downloaded & verified DOCX: {len(doc.paragraphs)} paragraphs, {len(doc.tables)} tables")
    if os.path.exists(temp_docx_path):
        os.remove(temp_docx_path)
        
    print("\n==================================================")
    print("🎉 ALL BACKEND & ANNEXURE TESTS PASSED!")
    print("==================================================")

if __name__ == '__main__':
    test_full_annexure_flow()
