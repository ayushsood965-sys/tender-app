import sys
import os
import json
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
from bs4 import BeautifulSoup, NavigableString, Tag
import zipfile

def set_cell_border(cell, **kwargs):
    """Set cell borders with custom line width, color, and style"""
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = f'w:{edge}'
            element = OxmlElement(tag)
            element.set(qn('w:val'), edge_data.get('val', 'single'))
            element.set(qn('w:sz'), str(edge_data.get('sz', 4)))
            element.set(qn('w:space'), str(edge_data.get('space', 0)))
            element.set(qn('w:color'), edge_data.get('color', '000000'))
            tcBorders.append(element)
        else:
            tag = f'w:{edge}'
            element = OxmlElement(tag)
            element.set(qn('w:val'), 'none')
            tcBorders.append(element)
    tcPr.append(tcBorders)

def set_cell_shading(cell, color_hex):
    """Set cell background shading"""
    shd_xml = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shd_xml)

def set_cell_margins(cell, top=80, bottom=80, left=120, right=120):
    """Set padding/margins inside a table cell (in twips)"""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
    tcPr.append(tcMar)

def add_p_bottom_border(p, color="CBD5E1", sz="6"):
    """Add a subtle bottom divider line under section headings"""
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="{sz}" w:space="2" w:color="{color}"/></w:pBdr>')
    pPr.append(pBdr)

def process_inline_nodes(p, elem, font_name="Montserrat", font_size=Pt(10.5), default_color=None):
    """Recursively process inline text formatting including bold, italics, underline, colors"""
    for child in elem.children:
        if isinstance(child, NavigableString):
            text = str(child)
            if text:
                r = p.add_run(text)
                r.font.name = font_name
                r.font.size = font_size
                if default_color:
                    r.font.color.rgb = default_color
        elif isinstance(child, Tag):
            text = child.get_text()
            if not text and child.name != 'br':
                continue

            if child.name == 'br':
                p.add_run('\n')
                continue

            r = p.add_run(text)
            r.font.name = font_name
            r.font.size = font_size

            # Bold
            if child.name in ['strong', 'b'] or 'bold' in child.get('style', '') or 'font-bold' in child.get('class', []):
                r.font.bold = True

            # Italic
            if child.name in ['em', 'i'] or 'italic' in child.get('style', ''):
                r.font.italic = True

            # Underline
            if child.name == 'u' or 'underline' in child.get('style', ''):
                r.font.underline = True

            # Custom colors
            style_str = child.get('style', '')
            if '#1e3a8a' in style_str or 'blue' in style_str:
                r.font.color.rgb = RGBColor(30, 58, 138)
            elif '#004821' in style_str or '#006400' in style_str or 'green' in style_str:
                r.font.color.rgb = RGBColor(0, 100, 0)
            elif '#475569' in style_str or '#334155' in style_str:
                r.font.color.rgb = RGBColor(71, 85, 105)
            elif default_color:
                r.font.color.rgb = default_color

def parse_html_elements_to_docx(soup, doc, logo_path=None, parent_align=None):
    """Parse HTML elements into python-docx paragraphs, callout boxes, and tables matching the web preview"""
    
    # Iterate through top-level block elements
    elements = soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'p', 'table', 'div', 'ul', 'ol', 'img'], recursive=False)
    
    # If no top-level recognized elements found, search children
    if not elements and soup.children:
        elements = [c for c in soup.children if isinstance(c, Tag)]

    for elem in elements:
        style_str = elem.get('style', '')

        # 1. Images (e.g. HPU Crest / Logo)
        if elem.name == 'img':
            if logo_path and os.path.exists(logo_path):
                p_img = doc.add_paragraph()
                p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p_img.paragraph_format.space_before = Pt(8)
                p_img.paragraph_format.space_after = Pt(12)
                p_img.add_run().add_picture(logo_path, width=Inches(1.6))
            continue

        # 2. Flex Containers (e.g. "No. 1-5/2025/HPU/SPS" on Left, "Dated: ..." on Right)
        if elem.name == 'div' and 'display: flex' in style_str:
            spans = elem.find_all(['span', 'div', 'p'])
            if len(spans) >= 2:
                t_flex = doc.add_table(rows=1, cols=2)
                t_flex.alignment = WD_TABLE_ALIGNMENT.CENTER
                row = t_flex.rows[0]
                row.cells[0].width = Inches(3.35)
                row.cells[1].width = Inches(3.35)
                
                # Left item
                p_left = row.cells[0].paragraphs[0]
                p_left.paragraph_format.space_before = Pt(2)
                p_left.paragraph_format.space_after = Pt(4)
                process_inline_nodes(p_left, spans[0], font_name="Montserrat", font_size=Pt(10))
                if len(p_left.runs) > 0:
                    p_left.runs[0].font.bold = True

                # Right item
                p_right = row.cells[1].paragraphs[0]
                p_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                p_right.paragraph_format.space_before = Pt(2)
                p_right.paragraph_format.space_after = Pt(4)
                process_inline_nodes(p_right, spans[1], font_name="Montserrat", font_size=Pt(10))
                if len(p_right.runs) > 0:
                    p_right.runs[0].font.bold = True
                continue

        # 3. Styled Callout / Placeholder Highlight Boxes
        if elem.name == 'div' and ('border-left' in style_str or 'background: #f8fafc' in style_str or 'dashed' in style_str or 'bg-slate-50' in elem.get('class', [])):
            t_box = doc.add_table(rows=1, cols=1)
            t_box.alignment = WD_TABLE_ALIGNMENT.CENTER
            cell = t_box.rows[0].cells[0]
            cell.width = Inches(6.7)
            
            is_dashed = 'dashed' in style_str
            set_cell_shading(cell, "F8FAFC")
            set_cell_margins(cell, top=120, bottom=120, left=160, right=160)
            
            if is_dashed:
                set_cell_border(cell,
                    top={'sz': 6, 'val': 'dashed', 'color': '94A3B8'},
                    bottom={'sz': 6, 'val': 'dashed', 'color': '94A3B8'},
                    right={'sz': 6, 'val': 'dashed', 'color': '94A3B8'},
                    left={'sz': 6, 'val': 'dashed', 'color': '94A3B8'}
                )
            else:
                set_cell_border(cell,
                    top={'sz': 4, 'val': 'single', 'color': 'CBD5E1'},
                    bottom={'sz': 4, 'val': 'single', 'color': 'CBD5E1'},
                    right={'sz': 4, 'val': 'single', 'color': 'CBD5E1'},
                    left={'sz': 24, 'val': 'single', 'color': '006400'} # Dark green 3pt left accent bar
                )

            cell_p = cell.paragraphs[0]
            if 'text-align: center' in style_str or 'center' in style_str:
                cell_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell_p.paragraph_format.space_before = Pt(4)
            cell_p.paragraph_format.space_after = Pt(4)
            cell_p.paragraph_format.line_spacing = 1.15

            inner_ps = elem.find_all('p')
            if inner_ps:
                for idx_p, ip in enumerate(inner_ps):
                    if idx_p > 0:
                        cell_p = cell.add_paragraph()
                        if 'text-align: center' in style_str or 'center' in style_str or 'text-align: center' in ip.get('style', ''):
                            cell_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        cell_p.paragraph_format.space_before = Pt(3)
                        cell_p.paragraph_format.space_after = Pt(3)
                        cell_p.paragraph_format.line_spacing = 1.15
                    process_inline_nodes(cell_p, ip, font_name="Montserrat", font_size=Pt(10.5))
            else:
                process_inline_nodes(cell_p, elem, font_name="Montserrat", font_size=Pt(10.5))
            continue

        # 4. Section & Document Headings
        if elem.name in ['h1', 'h2', 'h3', 'h4', 'h5']:
            p = doc.add_paragraph()
            align = WD_ALIGN_PARAGRAPH.LEFT
            if 'text-align: center' in style_str or 'center' in style_str or parent_align == 'center':
                align = WD_ALIGN_PARAGRAPH.CENTER
            elif 'text-align: right' in style_str or 'right' in style_str or parent_align == 'right':
                align = WD_ALIGN_PARAGRAPH.RIGHT

            p.alignment = align
            p.paragraph_format.keep_with_next = True

            if elem.name == 'h1':
                p.paragraph_format.space_before = Pt(14)
                p.paragraph_format.space_after = Pt(6)
                process_inline_nodes(p, elem, font_name="Montserrat", font_size=Pt(15), default_color=RGBColor(15, 23, 42))
            elif elem.name == 'h2':
                p.paragraph_format.space_before = Pt(10)
                p.paragraph_format.space_after = Pt(4)
                process_inline_nodes(p, elem, font_name="Montserrat", font_size=Pt(13.5), default_color=RGBColor(15, 23, 42))
            elif elem.name == 'h3':
                p.paragraph_format.space_before = Pt(9)
                p.paragraph_format.space_after = Pt(3)
                process_inline_nodes(p, elem, font_name="Montserrat", font_size=Pt(12), default_color=RGBColor(15, 23, 42))
            elif elem.name == 'h4': # Standard Section Header with underline/border
                p.paragraph_format.space_before = Pt(10)
                p.paragraph_format.space_after = Pt(3)
                add_p_bottom_border(p, color="CBD5E1", sz="6")
                process_inline_nodes(p, elem, font_name="Montserrat", font_size=Pt(11), default_color=RGBColor(15, 23, 42))
            else:
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after = Pt(2)
                process_inline_nodes(p, elem, font_name="Montserrat", font_size=Pt(10.5))

            if len(p.runs) > 0:
                p.runs[0].font.bold = True
            continue

        # 5. Paragraphs & Clauses
        if elem.name == 'p':
            text = elem.get_text().strip()
            # Check for embedded image
            img = elem.find('img')
            if img and logo_path and os.path.exists(logo_path):
                p_logo = doc.add_paragraph()
                p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p_logo.paragraph_format.space_before = Pt(8)
                p_logo.paragraph_format.space_after = Pt(12)
                p_logo.add_run().add_picture(logo_path, width=Inches(1.6))
                continue

            if not text:
                continue

            p = doc.add_paragraph()
            align = WD_ALIGN_PARAGRAPH.LEFT
            if 'text-align: center' in style_str or parent_align == 'center':
                align = WD_ALIGN_PARAGRAPH.CENTER
            elif 'text-align: right' in style_str or parent_align == 'right':
                align = WD_ALIGN_PARAGRAPH.RIGHT
            elif 'text-align: justify' in style_str or parent_align == 'justify' or len(text) > 80:
                align = WD_ALIGN_PARAGRAPH.JUSTIFY

            p.alignment = align
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(3.5)
            p.paragraph_format.line_spacing = 1.15
            process_inline_nodes(p, elem, font_name="Montserrat", font_size=Pt(10.5))
            continue

        # 6. Tables (Summary, BOQ, Compliance Matrix, Key-Values)
        if elem.name == 'table':
            trs = elem.find_all('tr')
            if not trs:
                continue

            num_rows = len(trs)
            num_cols = max(len(tr.find_all(['th', 'td'])) for tr in trs)
            table = doc.add_table(rows=num_rows, cols=num_cols)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER

            # Column width distribution matching A4 printable box (6.7 inches)
            total_width = Inches(6.7)
            if num_cols == 7: # Annexure-A Financial BOQ schedule
                col_widths = [Inches(0.5), Inches(1.6), Inches(1.8), Inches(0.6), Inches(0.8), Inches(0.6), Inches(0.8)]
            elif num_cols == 4: # Technical Compliance matrix
                col_widths = [Inches(0.6), Inches(3.0), Inches(2.0), Inches(1.1)]
            elif num_cols == 3: # Summary table
                col_widths = [Inches(0.7), Inches(2.2), Inches(3.8)]
            elif num_cols == 2: # Two-column key-value / bidder details
                col_widths = [Inches(2.7), Inches(4.0)]
            else:
                col_widths = [total_width / num_cols] * num_cols

            for r_idx, tr in enumerate(trs):
                cells = tr.find_all(['th', 'td'])
                row = table.rows[r_idx]
                is_header = (r_idx == 0 and tr.find('th') is not None)

                for c_idx, cell_data in enumerate(cells):
                    if c_idx >= num_cols:
                        break
                    cell = row.cells[c_idx]
                    cell.width = col_widths[c_idx]

                    set_cell_border(cell,
                        top={'sz': 4, 'val': 'single', 'color': '000000'},
                        bottom={'sz': 4, 'val': 'single', 'color': '000000'},
                        left={'sz': 4, 'val': 'single', 'color': '000000'},
                        right={'sz': 4, 'val': 'single', 'color': '000000'}
                    )

                    set_cell_margins(cell, top=60, bottom=60, left=100, right=100)

                    if is_header or cell_data.name == 'th':
                        set_cell_shading(cell, "F0F0F0")

                    p = cell.paragraphs[0]
                    p.paragraph_format.space_before = Pt(2)
                    p.paragraph_format.space_after = Pt(2)
                    p.paragraph_format.line_spacing = 1.1

                    process_inline_nodes(p, cell_data, font_name="Montserrat", font_size=Pt(9.5))
                    if is_header or cell_data.name == 'th':
                        for r in p.runs:
                            r.font.bold = True

            # Add a small spacing paragraph after table
            p_space = doc.add_paragraph()
            p_space.paragraph_format.space_before = Pt(1)
            p_space.paragraph_format.space_after = Pt(4)
            continue

        # 7. Divs & Containers (Recurse with inherited styling)
        if elem.name == 'div':
            align = None
            if 'text-align: right' in style_str:
                align = 'right'
            elif 'text-align: center' in style_str:
                align = 'center'
            elif 'text-align: justify' in style_str:
                align = 'justify'
            else:
                align = parent_align

            # Check if div contains only text nodes directly
            direct_text = "".join([str(c) for c in elem.children if isinstance(c, NavigableString)]).strip()
            has_child_blocks = elem.find(['h1', 'h2', 'h3', 'h4', 'h5', 'p', 'table', 'div', 'ul', 'ol']) is not None
            
            if direct_text and not has_child_blocks:
                p = doc.add_paragraph()
                if align == 'right':
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                elif align == 'center':
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                elif align == 'justify':
                    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after = Pt(3.5)
                p.paragraph_format.line_spacing = 1.15
                process_inline_nodes(p, elem, font_name="Montserrat", font_size=Pt(10.5))
            else:
                parse_html_elements_to_docx(elem, doc, logo_path, parent_align=align)

def is_annexure_page(html):
    """Detect if the page is a discrete Annexure proforma/letter that requires its own fresh page"""
    html_low = html.lower()
    soup = BeautifulSoup(html[:1000], 'html.parser')
    for heading in soup.find_all(['h1', 'h2', 'h3']):
        if 'annexure' in heading.get_text().lower():
            return True
    if 'annexure-' in html_low[:300] or 'annexure -' in html_low[:300] or 'annexure ' in html_low[:150]:
        return True
    return False

def is_terms_continuation(html, prev_was_terms):
    """Detect if the page is a continuation of the Terms & Conditions stream"""
    soup = BeautifulSoup(html[:400], 'html.parser')
    text = soup.get_text().strip()
    
    if is_annexure_page(html):
        return False
    
    if prev_was_terms:
        if 'section ' in text.lower()[:300] or text[:20].strip().startswith(('1.', '2.', '3.', '4.', '5.', '6.', '7.', '8.', '9.', '10.', '11.', '12.', '13.', '14.', '15.', '16.', '17.', '18.', '19.', '20.', '21.', '22.', '23.', '24.', '25.', '26.')):
            return True
    return False

def generate_docx_from_json(input_json_path, output_docx_path, logo_path, watermark_path):
    with open(input_json_path, 'r', encoding='utf-8') as f:
        data = json.load(f)

    pages = data.get("pages", [])
    doc = docx.Document()
    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.85)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    section.header_distance = Inches(0.3)
    section.footer_distance = Inches(0.55)

    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Montserrat'
    normal_style.font.size = Pt(10.5)
    normal_style.font.color.rgb = RGBColor(0, 0, 0)

    # Process each page from the preview document with smart continuous terms flow
    is_terms = False
    for idx, page_html in enumerate(pages):
        annex = is_annexure_page(page_html)
        cont = is_terms_continuation(page_html, is_terms)

        if idx == 0:
            needs_break = False
        elif annex:
            needs_break = True
            is_terms = False
        elif cont:
            needs_break = False # Continuous terms flow in Word prevents artificial blank pages!
            is_terms = True
        else:
            needs_break = True
            if 'section 1:' in page_html.lower() or 'section 1 :' in page_html.lower() or 'section 1' in page_html.lower()[:300]:
                is_terms = True
            else:
                is_terms = False

        if needs_break:
            doc.add_page_break()

        soup = BeautifulSoup(page_html, 'html.parser')
        parse_html_elements_to_docx(soup, doc, logo_path)

    temp_raw = output_docx_path + ".temp"
    doc.save(temp_raw)

    # XML Templates for Watermark & Footers with Dynamic Page Numbers
    header_xml = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:hdr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships" xmlns:v="urn:schemas-microsoft-com:vml" xmlns:o="urn:schemas-microsoft-com:office:office" xmlns:w10="urn:schemas-microsoft-com:office:word">
    <w:p>
        <w:pPr><w:pStyle w:val="Header"/></w:pPr>
        <w:r>
            <w:pict>
                <v:shapetype id="_x0000_t75" coordsize="21600,21600" o:spt="75" o:preferrelative="t" path="m@4@5l@4@11@9@11@9@5xe" filled="f" stroked="f">
                    <v:stroke joinstyle="miter"/>
                    <v:formulas>
                        <v:f eqn="if lineDrawn pixelLineWidth 0"/>
                        <v:f eqn="sum @0 1 0"/>
                        <v:f eqn="sum 0 0 @1"/>
                        <v:f eqn="prod @2 1 2"/>
                        <v:f eqn="prod @3 21600 pixelWidth"/>
                        <v:f eqn="prod @3 21600 pixelHeight"/>
                        <v:f eqn="sum @0 0 1"/>
                        <v:f eqn="prod @6 1 2"/>
                        <v:f eqn="prod @7 21600 pixelWidth"/>
                        <v:f eqn="sum @8 21600 0"/>
                        <v:f eqn="prod @7 21600 pixelHeight"/>
                        <v:f eqn="sum @10 21600 0"/>
                    </v:formulas>
                    <v:path o:extrusionok="f" gradientshapeok="t" o:connecttype="rect"/>
                    <o:lock v:ext="edit" aspectratio="t"/>
                </v:shapetype>
                <v:shape id="WordPictureWatermark" o:spid="_x0000_s1029" type="#_x0000_t75" style="position:absolute;margin-left:0;margin-top:0;width:480pt;height:360pt;z-index:-251657216;mso-position-horizontal:center;mso-position-horizontal-relative:margin;mso-position-vertical:center;mso-position-vertical-relative:margin" o:allowincell="f">
                    <v:imagedata r:id="rId1" o:title="HPU Watermark" gain="19661f" blacklevel="22938f"/>
                    <w10:wrap anchorx="margin" anchory="margin"/>
                </v:shape>
            </w:pict>
        </w:r>
    </w:p>
</w:hdr>"""

    header_rels = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
    <Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/image" Target="media/watermark.png"/>
</Relationships>"""

    footer_xml = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:ftr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
    <w:p>
        <w:pPr>
            <w:pStyle w:val="Footer"/>
            <w:spacing w:before="0" w:after="0" w:line="240" w:lineRule="auto"/>
            <w:jc w:val="right"/>
        </w:pPr>
        <w:r>
            <w:rPr>
                <w:rFonts w:ascii="Montserrat" w:hAnsi="Montserrat"/>
                <w:b/>
                <w:sz w:val="20"/>
                <w:color w:val="1E293B"/>
            </w:rPr>
            <w:t xml:space="preserve">Page </w:t>
        </w:r>
        <w:r>
            <w:rPr>
                <w:rFonts w:ascii="Montserrat" w:hAnsi="Montserrat"/>
                <w:b/>
                <w:sz w:val="20"/>
                <w:color w:val="1E293B"/>
            </w:rPr>
            <w:fldChar w:fldCharType="begin"/>
        </w:r>
        <w:r>
            <w:rPr>
                <w:rFonts w:ascii="Montserrat" w:hAnsi="Montserrat"/>
                <w:b/>
                <w:sz w:val="20"/>
                <w:color w:val="1E293B"/>
            </w:rPr>
            <w:instrText xml:space="preserve"> PAGE </w:instrText>
        </w:r>
        <w:r>
            <w:rPr>
                <w:rFonts w:ascii="Montserrat" w:hAnsi="Montserrat"/>
                <w:b/>
                <w:sz w:val="20"/>
                <w:color w:val="1E293B"/>
            </w:rPr>
            <w:fldChar w:fldCharType="end"/>
        </w:r>
        <w:r>
            <w:rPr>
                <w:rFonts w:ascii="Montserrat" w:hAnsi="Montserrat"/>
                <w:b/>
                <w:sz w:val="20"/>
                <w:color w:val="1E293B"/>
            </w:rPr>
            <w:t xml:space="preserve"> of </w:t>
        </w:r>
        <w:r>
            <w:rPr>
                <w:rFonts w:ascii="Montserrat" w:hAnsi="Montserrat"/>
                <w:b/>
                <w:sz w:val="20"/>
                <w:color w:val="1E293B"/>
            </w:rPr>
            <w:fldChar w:fldCharType="begin"/>
        </w:r>
        <w:r>
            <w:rPr>
                <w:rFonts w:ascii="Montserrat" w:hAnsi="Montserrat"/>
                <w:b/>
                <w:sz w:val="20"/>
                <w:color w:val="1E293B"/>
            </w:rPr>
            <w:instrText xml:space="preserve"> NUMPAGES </w:instrText>
        </w:r>
        <w:r>
            <w:rPr>
                <w:rFonts w:ascii="Montserrat" w:hAnsi="Montserrat"/>
                <w:b/>
                <w:sz w:val="20"/>
                <w:color w:val="1E293B"/>
            </w:rPr>
            <w:fldChar w:fldCharType="end"/>
        </w:r>
    </w:p>
</w:ftr>"""

    watermark_bytes = b""
    if os.path.exists(watermark_path):
        with open(watermark_path, 'rb') as f:
            watermark_bytes = f.read()

    with zipfile.ZipFile(temp_raw, 'r') as zin:
        with zipfile.ZipFile(output_docx_path, 'w') as zout:
            for item in zin.infolist():
                content = zin.read(item.filename)
                
                if item.filename == '[Content_Types].xml':
                    ct_str = content.decode('utf-8')
                    injections = ""
                    if 'Extension="png"' not in ct_str:
                        injections += '<Default Extension="png" ContentType="image/png"/>'
                    if 'header1.xml' not in ct_str:
                        injections += '<Override PartName="/word/header1.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.header+xml"/><Override PartName="/word/footer1.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.footer+xml"/>'
                    if injections:
                        ct_str = ct_str.replace('</Types>', f'{injections}</Types>')
                    zout.writestr(item, ct_str.encode('utf-8'))

                elif item.filename == 'word/_rels/document.xml.rels':
                    rels_str = content.decode('utf-8')
                    if 'header1.xml' not in rels_str:
                        rels_str = rels_str.replace('</Relationships>', '<Relationship Id="rIdHeader1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/header" Target="header1.xml"/><Relationship Id="rIdFooter1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/footer" Target="footer1.xml"/></Relationships>')
                    zout.writestr(item, rels_str.encode('utf-8'))

                elif item.filename == 'word/document.xml':
                    doc_str = content.decode('utf-8')
                    pg_borders_xml = '<w:headerReference w:type="default" r:id="rIdHeader1"/><w:footerReference w:type="default" r:id="rIdFooter1"/><w:pgBorders w:offsetFrom="page"><w:top w:val="thinThickThinSmallGap" w:sz="24" w:space="20" w:color="006400"/><w:left w:val="thinThickThinSmallGap" w:sz="24" w:space="20" w:color="006400"/><w:bottom w:val="thinThickThinSmallGap" w:sz="24" w:space="20" w:color="006400"/><w:right w:val="thinThickThinSmallGap" w:sz="24" w:space="20" w:color="006400"/></w:pgBorders>'
                    doc_str = doc_str.replace('</w:sectPr>', f'{pg_borders_xml}</w:sectPr>')
                    zout.writestr(item, doc_str.encode('utf-8'))

                else:
                    zout.writestr(item, content)

            zout.writestr('word/header1.xml', header_xml.encode('utf-8'))
            zout.writestr('word/_rels/header1.xml.rels', header_rels.encode('utf-8'))
            zout.writestr('word/footer1.xml', footer_xml.encode('utf-8'))
            if watermark_bytes:
                zout.writestr('word/media/watermark.png', watermark_bytes)

    if os.path.exists(temp_raw):
        os.remove(temp_raw)

    print(f"Generated complete styled docx -> {output_docx_path}")

if __name__ == '__main__':
    if len(sys.argv) < 3:
        print("Usage: python docx_generator.py <input_json> <output_docx>")
        sys.exit(1)
    
    in_json = sys.argv[1]
    out_docx = sys.argv[2]
    script_dir = os.path.dirname(os.path.abspath(__file__))
    
    logo = os.path.join(script_dir, "assets", "hpu_logo.png")
    if not os.path.exists(logo):
        logo = os.path.join(script_dir, "..", "client", "public", "hpu_logo.png")

    watermark = os.path.join(script_dir, "assets", "hpu_watermark.png")
    if not os.path.exists(watermark):
        watermark = os.path.join(script_dir, "..", "client", "public", "hpu_watermark.png")

    generate_docx_from_json(in_json, out_docx, logo, watermark)
