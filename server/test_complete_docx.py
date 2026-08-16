import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
from bs4 import BeautifulSoup
import os
import zipfile
import shutil
import json

def generate_perfect_docx(pages_html_list, output_docx_path, logo_path, watermark_path):
    # 1. Start with python-docx Document
    doc = docx.Document()

    # Configure Section: A4, 0.8 inch margins
    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    section.header_distance = Inches(0.4)
    section.footer_distance = Inches(0.4)

    # Style: Montserrat Font
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Montserrat'
    normal_style.font.size = Pt(10.5)
    normal_style.font.color.rgb = RGBColor(0, 0, 0)

    for idx, page_html in enumerate(pages_html_list):
        soup = BeautifulSoup(page_html, 'html.parser')
        is_cover = (idx == 0)
        
        # Parse Page Content
        if is_cover:
            # Cover Page
            p_top = doc.add_paragraph()
            p_top.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_top.paragraph_format.space_before = Pt(30)
            p_top.paragraph_format.space_after = Pt(3)
            r = p_top.add_run("Himachal Pradesh University,")
            r.font.name = "Montserrat"
            r.font.size = Pt(16)
            r.font.bold = True

            p_sub = doc.add_paragraph()
            p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_sub.paragraph_format.space_before = Pt(0)
            p_sub.paragraph_format.space_after = Pt(3)
            r = p_sub.add_run("(NAAC Accredited 'A' Grade University)")
            r.font.name = "Montserrat"
            r.font.size = Pt(13)
            r.font.bold = True

            # Find department text
            dept_text = ""
            for h in soup.find_all(['h3', 'h2', 'p']):
                t = h.get_text().strip()
                if "Department" in t or "Centre" in t or "CDOE" in t:
                    dept_text = t
                    break
            
            if dept_text:
                p_dept = doc.add_paragraph()
                p_dept.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p_dept.paragraph_format.space_before = Pt(0)
                p_dept.paragraph_format.space_after = Pt(16)
                r = p_dept.add_run(dept_text)
                r.font.name = "Montserrat"
                r.font.size = Pt(13)
                r.font.bold = True

            # Logo
            if logo_path and os.path.exists(logo_path):
                p_logo = doc.add_paragraph()
                p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p_logo.paragraph_format.space_before = Pt(10)
                p_logo.paragraph_format.space_after = Pt(16)
                p_logo.add_run().add_picture(logo_path, width=Inches(1.8))

            # Title
            h1 = soup.find('h1')
            title_text = h1.get_text().strip() if h1 else "Bid Document"
            p_title = doc.add_paragraph()
            p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_title.paragraph_format.space_before = Pt(16)
            p_title.paragraph_format.space_after = Pt(18)
            r = p_title.add_run(title_text)
            r.font.name = "Montserrat"
            r.font.size = Pt(15)
            r.font.bold = True

            # Contact
            p_web = doc.add_paragraph()
            p_web.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_web.paragraph_format.space_before = Pt(16)
            p_web.paragraph_format.space_after = Pt(2)
            r = p_web.add_run("Website: https://www.hpuniv.ac.in")
            r.font.name = "Montserrat"
            r.font.size = Pt(11)
            r.font.bold = True

            p_mail = doc.add_paragraph()
            p_mail.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_mail.paragraph_format.space_before = Pt(0)
            p_mail.paragraph_format.space_after = Pt(6)
            r = p_mail.add_run("E-mail: info@hpuniv.ac.in")
            r.font.name = "Montserrat"
            r.font.size = Pt(11)
            r.font.bold = True

        else:
            # Regular Content Page
            for elem in soup.find_all(['h1', 'h2', 'h3', 'h4', 'p', 'table', 'ul', 'ol']):
                if elem.name in ['h1', 'h2', 'h3', 'h4']:
                    p = doc.add_paragraph()
                    align = WD_ALIGN_PARAGRAPH.LEFT
                    if 'center' in elem.get('style', '') or 'center' in elem.get('align', ''):
                        align = WD_ALIGN_PARAGRAPH.CENTER
                    p.alignment = align
                    p.paragraph_format.space_before = Pt(12)
                    p.paragraph_format.space_after = Pt(6)
                    p.paragraph_format.keep_with_next = True
                    r = p.add_run(elem.get_text().strip())
                    r.font.name = "Montserrat"
                    r.font.bold = True
                    r.font.size = Pt(12)

                elif elem.name == 'p':
                    text = elem.get_text().strip()
                    if not text:
                        continue
                    p = doc.add_paragraph()
                    p.paragraph_format.space_before = Pt(3)
                    p.paragraph_format.space_after = Pt(5)
                    p.paragraph_format.line_spacing = 1.15
                    for child in elem.children:
                        if isinstance(child, str):
                            r = p.add_run(child)
                            r.font.name = "Montserrat"
                            r.font.size = Pt(10.5)
                        elif child.name in ['strong', 'b']:
                            r = p.add_run(child.get_text())
                            r.font.name = "Montserrat"
                            r.font.size = Pt(10.5)
                            r.font.bold = True
                        elif child.name in ['em', 'i']:
                            r = p.add_run(child.get_text())
                            r.font.name = "Montserrat"
                            r.font.size = Pt(10.5)
                            r.font.italic = True
                        else:
                            r = p.add_run(child.get_text())
                            r.font.name = "Montserrat"
                            r.font.size = Pt(10.5)

                elif elem.name == 'table':
                    trs = elem.find_all('tr')
                    if trs:
                        num_rows = len(trs)
                        num_cols = max(len(tr.find_all(['th', 'td'])) for tr in trs)
                        table = doc.add_table(rows=num_rows, cols=num_cols)
                        table.alignment = WD_TABLE_ALIGNMENT.CENTER
                        
                        col_widths = [Inches(6.6 / num_cols)] * num_cols
                        if num_cols == 3:
                            col_widths = [Inches(0.8), Inches(2.2), Inches(3.6)]

                        for r_idx, tr in enumerate(trs):
                            cells = tr.find_all(['th', 'td'])
                            row = table.rows[r_idx]
                            is_header = (r_idx == 0 and tr.find('th') is not None)

                            for c_idx, cell_data in enumerate(cells):
                                if c_idx >= num_cols:
                                    break
                                cell = row.cells[c_idx]
                                cell.width = col_widths[c_idx]

                                # Borders
                                tcPr = cell._tc.get_or_add_tcPr()
                                borders_xml = parse_xml(f'<w:tcBorders {nsdecls("w")}><w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/><w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/><w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/><w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/></w:tcBorders>')
                                tcPr.append(borders_xml)

                                if is_header:
                                    shd_xml = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F0F0F0"/>')
                                    tcPr.append(shd_xml)

                                p = cell.paragraphs[0]
                                p.paragraph_format.space_before = Pt(3)
                                p.paragraph_format.space_after = Pt(3)
                                r = p.add_run(cell_data.get_text().strip())
                                r.font.name = "Montserrat"
                                r.font.size = Pt(9.5)
                                if is_header or cell_data.name == 'th' or cell_data.find(['strong', 'b']):
                                    r.font.bold = True

        # Add Page Break between pages
        if idx < len(pages_html_list) - 1:
            doc.add_page_break()

    # Save to temp file
    temp_raw = "temp_raw_doc.docx"
    doc.save(temp_raw)

    # 2. Inject XML Enhancements (Triple Green Page Border + Watermark Shape + Page Numbers Footer)
    # Prepare Watermark VML Header and Footer XML
    header_xml = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:hdr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships" xmlns:v="urn:schemas-microsoft-com:vml" xmlns:o="urn:schemas-microsoft-com:office:office" xmlns:w10="urn:schemas-microsoft-com:office:word">
    <w:p>
        <w:pPr><w:pStyle w:val="Header"/></w:pPr>
        <w:r>
            <w:pict>
                <v:shapetype id="_x0000_t75" coordsize="21600,21600" o:spt="75" o:preferrelative="t" path="m@4@5l@4@11@9@11@9@5xe" filled="f" stroked="f">
                    <v:stroke joinstyle="miter"/>
                    <v:formulas>
                        <v:f eqn="if lineDrawn pixelLineWidth 0"/>
                        <v:f eqn="sum @0 1 0"/>
                        <v:f eqn="sum 0 0 @1"/>
                        <v:f eqn="prod @2 1 2"/>
                        <v:f eqn="prod @3 21600 pixelWidth"/>
                        <v:f eqn="prod @3 21600 pixelHeight"/>
                        <v:f eqn="sum @0 0 1"/>
                        <v:f eqn="prod @6 1 2"/>
                        <v:f eqn="prod @7 21600 pixelWidth"/>
                        <v:f eqn="sum @8 21600 0"/>
                        <v:f eqn="prod @7 21600 pixelHeight"/>
                        <v:f eqn="sum @10 21600 0"/>
                    </v:formulas>
                    <v:path o:extrusionok="f" gradientshapeok="t" o:connecttype="rect"/>
                    <o:lock v:ext="edit" aspectratio="t"/>
                </v:shapetype>
                <v:shape id="WordPictureWatermark" o:spid="_x0000_s1029" type="#_x0000_t75" style="position:absolute;margin-left:0;margin-top:0;width:480pt;height:360pt;z-index:-251657216;mso-position-horizontal:center;mso-position-horizontal-relative:margin;mso-position-vertical:center;mso-position-vertical-relative:margin" o:allowincell="f">
                    <v:imagedata r:id="rId1" o:title="HPU Watermark" gain="19661f" blacklevel="22938f"/>
                    <w10:wrap anchorx="margin" anchory="margin"/>
                </v:shape>
            </w:pict>
        </w:r>
    </w:p>
</w:hdr>"""

    header_rels = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
    <Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/image" Target="media/watermark.png"/>
</Relationships>"""

    footer_xml = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:ftr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
    <w:p>
        <w:pPr>
            <w:pStyle w:val="Footer"/>
            <w:jc w:val="right"/>
        </w:pPr>
        <w:r>
            <w:rPr>
                <w:rFonts w:ascii="Montserrat" w:hAnsi="Montserrat"/>
                <w:b/>
                <w:sz w:val="20"/>
            </w:rPr>
            <w:fldChar w:fldCharType="begin"/>
        </w:r>
        <w:r>
            <w:rPr>
                <w:rFonts w:ascii="Montserrat" w:hAnsi="Montserrat"/>
                <w:b/>
                <w:sz w:val="20"/>
            </w:rPr>
            <w:instrText xml:space="preserve"> PAGE </w:instrText>
        </w:r>
        <w:r>
            <w:rPr>
                <w:rFonts w:ascii="Montserrat" w:hAnsi="Montserrat"/>
                <w:b/>
                <w:sz w:val="20"/>
            </w:rPr>
            <w:fldChar w:fldCharType="separate"/>
        </w:r>
        <w:r>
            <w:rPr>
                <w:rFonts w:ascii="Montserrat" w:hAnsi="Montserrat"/>
                <w:b/>
                <w:sz w:val="20"/>
            </w:rPr>
            <w:fldChar w:fldCharType="end"/>
        </w:r>
        <w:r>
            <w:rPr>
                <w:rFonts w:ascii="Montserrat" w:hAnsi="Montserrat"/>
                <w:b/>
                <w:sz w:val="20"/>
            </w:rPr>
            <w:t>/</w:t>
        </w:r>
        <w:r>
            <w:rPr>
                <w:rFonts w:ascii="Montserrat" w:hAnsi="Montserrat"/>
                <w:b/>
                <w:sz w:val="20"/>
            </w:rPr>
            <w:fldChar w:fldCharType="begin"/>
        </w:r>
        <w:r>
            <w:rPr>
                <w:rFonts w:ascii="Montserrat" w:hAnsi="Montserrat"/>
                <w:b/>
                <w:sz w:val="20"/>
            </w:rPr>
            <w:instrText xml:space="preserve"> NUMPAGES </w:instrText>
        </w:r>
        <w:r>
            <w:rPr>
                <w:rFonts w:ascii="Montserrat" w:hAnsi="Montserrat"/>
                <w:b/>
                <w:sz w:val="20"/>
            </w:rPr>
            <w:fldChar w:fldCharType="separate"/>
        </w:r>
        <w:r>
            <w:rPr>
                <w:rFonts w:ascii="Montserrat" w:hAnsi="Montserrat"/>
                <w:b/>
                <w:sz w:val="20"/>
            </w:rPr>
            <w:fldChar w:fldCharType="end"/>
        </w:r>
    </w:p>
</w:ftr>"""

    # Read watermark image binary
    with open(watermark_path, 'rb') as f:
        watermark_bytes = f.read()

    # Re-package zip with injected watermark, footer, and triple green borders
    with zipfile.ZipFile(temp_raw, 'r') as zin:
        with zipfile.ZipFile(output_docx_path, 'w') as zout:
            for item in zin.infolist():
                data = zin.read(item.filename)
                
                # 1. Update [Content_Types].xml to register header and footer
                if item.filename == '[Content_Types].xml':
                    ct_str = data.decode('utf-8')
                    if 'header1.xml' not in ct_str:
                        ct_str = ct_str.replace('</Types>', '<Override PartName="/word/header1.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.header+xml"/><Override PartName="/word/footer1.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.footer+xml"/></Types>')
                    zout.writestr(item, ct_str.encode('utf-8'))

                # 2. Update word/_rels/document.xml.rels to link header1 and footer1
                elif item.filename == 'word/_rels/document.xml.rels':
                    rels_str = data.decode('utf-8')
                    if 'header1.xml' not in rels_str:
                        rels_str = rels_str.replace('</Relationships>', '<Relationship Id="rIdHeader1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/header" Target="header1.xml"/><Relationship Id="rIdFooter1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/footer" Target="footer1.xml"/></Relationships>')
                    zout.writestr(item, rels_str.encode('utf-8'))

                # 3. Update word/document.xml to include headerReference, footerReference, and pgBorders
                elif item.filename == 'word/document.xml':
                    doc_str = data.decode('utf-8')
                    # Inject header & footer reference and pgBorders
                    pg_borders_xml = '<w:headerReference w:type="default" r:id="rIdHeader1"/><w:footerReference w:type="default" r:id="rIdFooter1"/><w:pgBorders w:offsetFrom="page"><w:top w:val="thinThickThinSmallGap" w:sz="24" w:space="24" w:color="004821"/><w:left w:val="thinThickThinSmallGap" w:sz="24" w:space="24" w:color="004821"/><w:bottom w:val="thinThickThinSmallGap" w:sz="24" w:space="24" w:color="004821"/><w:right w:val="thinThickThinSmallGap" w:sz="24" w:space="24" w:color="004821"/></w:pgBorders>'
                    doc_str = doc_str.replace('</w:sectPr>', f'{pg_borders_xml}</w:sectPr>')
                    zout.writestr(item, doc_str.encode('utf-8'))

                else:
                    zout.writestr(item, data)

            # Write header1.xml, header1.xml.rels, footer1.xml, and media/watermark.png
            zout.writestr('word/header1.xml', header_xml.encode('utf-8'))
            zout.writestr('word/_rels/header1.xml.rels', header_rels.encode('utf-8'))
            zout.writestr('word/footer1.xml', footer_xml.encode('utf-8'))
            zout.writestr('word/media/watermark.png', watermark_bytes)

    if os.path.exists(temp_raw):
        os.remove(temp_raw)

    print(f"Generated perfect DOCX with triple borders, watermark, and footer at {output_docx_path}")

if __name__ == '__main__':
    sample_page_1 = """<div><h1>Bid Document for High Resolution XRD</h1></div>"""
    sample_page_2 = """<div><h3>SUMMARY</h3><table><tr><th>Sr.</th><th>Description</th><th>Details</th></tr><tr><td>1</td><td>Item</td><td>XRD</td></tr></table><p><strong>1.1 General Terms:</strong> All conditions must be adhered to.</p></div>"""
    generate_perfect_docx([sample_page_1, sample_page_2], "test_styled_xrd.docx", "assets/hpu_logo.png", "assets/hpu_watermark.png")
