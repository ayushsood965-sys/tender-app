import docx
from docx.shared import Inches, Pt, RGBColor, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
from bs4 import BeautifulSoup
import os
import re
import zipfile

def set_cell_border(cell, **kwargs):
    """
    Set cell borders
    kwargs: top, bottom, left, right
    values: dict(sz=12, val='single', color='000000', space='0')
    """
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = f'w:{edge}'
            element = OxmlElement(tag)
            element.set(qn('w:val'), edge_data.get('val', 'single'))
            element.set(qn('w:sz'), str(edge_data.get('sz', 4)))
            element.set(qn('w:space'), str(edge_data.get('space', 0)))
            element.set(qn('w:color'), edge_data.get('color', '000000'))
            tcBorders.append(element)
    tcPr.append(tcBorders)

def set_cell_shading(cell, color_hex):
    shading_xml = f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>'
    cell._tc.get_or_add_tcPr().append(parse_xml(shading_xml))

def parse_html_page_to_docx(soup, doc, is_cover=False, logo_path=None):
    """
    Convert a single HTML page into styled docx paragraphs and tables
    """
    if is_cover:
        # Cover page layout
        p_top = doc.add_paragraph()
        p_top.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_top.paragraph_format.space_before = Pt(20)
        p_top.paragraph_format.space_after = Pt(2)
        r = p_top.add_run("Himachal Pradesh University,")
        r.font.name = "Montserrat"
        r.font.size = Pt(15)
        r.font.bold = True

        p_sub = doc.add_paragraph()
        p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_sub.paragraph_format.space_before = Pt(0)
        p_sub.paragraph_format.space_after = Pt(2)
        r = p_sub.add_run("(NAAC Accredited 'A' Grade University)")
        r.font.name = "Montserrat"
        r.font.size = Pt(13)
        r.font.bold = True

        # Extract Department Name
        h3s = soup.find_all(['h3', 'h2'])
        dept_name = ""
        for h in h3s:
            t = h.get_text().strip()
            if t and not "Himachal" in t and not "NAAC" in t:
                dept_name = t
                break
        
        if dept_name:
            p_dept = doc.add_paragraph()
            p_dept.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_dept.paragraph_format.space_before = Pt(0)
            p_dept.paragraph_format.space_after = Pt(14)
            r = p_dept.add_run(dept_name)
            r.font.name = "Montserrat"
            r.font.size = Pt(13)
            r.font.bold = True

        # Add Logo
        if logo_path and os.path.exists(logo_path):
            p_logo = doc.add_paragraph()
            p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_logo.paragraph_format.space_before = Pt(10)
            p_logo.paragraph_format.space_after = Pt(16)
            p_logo.add_run().add_picture(logo_path, width=Inches(1.8))

        # Title
        h1 = soup.find('h1')
        title_text = h1.get_text().strip() if h1 else "Bid Document"
        
        p_title = doc.add_paragraph()
        p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_title.paragraph_format.space_before = Pt(14)
        p_title.paragraph_format.space_after = Pt(16)
        r = p_title.add_run(title_text)
        r.font.name = "Montserrat"
        r.font.size = Pt(15)
        r.font.bold = True

        # Website and Email
        website_text = "Website: https://www.hpuniv.ac.in"
        email_text = ""
        for p in soup.find_all('p'):
            t = p.get_text().strip()
            if "Website:" in t:
                website_text = t
            elif "E-mail:" in t or "Email:" in t:
                email_text = t

        p_web = doc.add_paragraph()
        p_web.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_web.paragraph_format.space_before = Pt(16)
        p_web.paragraph_format.space_after = Pt(2)
        r = p_web.add_run(website_text)
        r.font.name = "Montserrat"
        r.font.size = Pt(11)
        r.font.bold = True

        if email_text:
            p_mail = doc.add_paragraph()
            p_mail.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_mail.paragraph_format.space_before = Pt(0)
            p_mail.paragraph_format.space_after = Pt(8)
            r = p_mail.add_run(email_text)
            r.font.name = "Montserrat"
            r.font.size = Pt(11)
            r.font.bold = True

        return

    # For Content / Terms Pages
    for elem in soup.find_all(['h1', 'h2', 'h3', 'h4', 'p', 'table', 'ul', 'ol'], recursive=False):
        if elem.name in ['h1', 'h2', 'h3', 'h4']:
            p = doc.add_paragraph()
            align = WD_ALIGN_PARAGRAPH.LEFT
            if 'text-align: center' in elem.get('style', '') or 'center' in elem.get('align', ''):
                align = WD_ALIGN_PARAGRAPH.CENTER
            p.alignment = align
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.keep_with_next = True

            text = elem.get_text().strip()
            r = p.add_run(text)
            r.font.name = "Montserrat"
            r.font.bold = True
            if elem.name in ['h1', 'h2']:
                r.font.size = Pt(13)
            else:
                r.font.size = Pt(12)

        elif elem.name == 'p':
            text = elem.get_text().strip()
            if not text:
                continue
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.line_spacing = 1.15

            # Process inline elements (strong, b, a, span)
            for child in elem.children:
                if isinstance(child, str):
                    r = p.add_run(child)
                    r.font.name = "Montserrat"
                    r.font.size = Pt(10.5)
                elif child.name in ['strong', 'b']:
                    r = p.add_run(child.get_text())
                    r.font.name = "Montserrat"
                    r.font.size = Pt(10.5)
                    r.font.bold = True
                elif child.name in ['em', 'i']:
                    r = p.add_run(child.get_text())
                    r.font.name = "Montserrat"
                    r.font.size = Pt(10.5)
                    r.font.italic = True
                else:
                    r = p.add_run(child.get_text())
                    r.font.name = "Montserrat"
                    r.font.size = Pt(10.5)

        elif elem.name == 'table':
            trs = elem.find_all('tr')
            if not trs:
                continue
            
            num_rows = len(trs)
            # determine cols
            num_cols = max(len(tr.find_all(['th', 'td'])) for tr in trs)
            
            table = doc.add_table(rows=num_rows, cols=num_cols)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.autofit = False

            # Set table width
            total_width = Inches(6.6)
            col_widths = [total_width / num_cols] * num_cols
            if num_cols == 3:
                col_widths = [Inches(0.8), Inches(2.2), Inches(3.6)]
            
            for row_idx, tr in enumerate(trs):
                cells = tr.find_all(['th', 'td'])
                row = table.rows[row_idx]
                is_header = (row_idx == 0 and tr.find('th') is not None)

                for col_idx, cell_data in enumerate(cells):
                    if col_idx >= num_cols:
                        break
                    cell = row.cells[col_idx]
                    cell.width = col_widths[col_idx]
                    
                    # Set border and padding
                    set_cell_border(cell, 
                        top={'sz': 4, 'val': 'single', 'color': '000000'},
                        bottom={'sz': 4, 'val': 'single', 'color': '000000'},
                        left={'sz': 4, 'val': 'single', 'color': '000000'},
                        right={'sz': 4, 'val': 'single', 'color': '000000'}
                    )
                    
                    if is_header:
                        set_cell_shading(cell, "F0F0F0")

                    p = cell.paragraphs[0]
                    p.paragraph_format.space_before = Pt(3)
                    p.paragraph_format.space_after = Pt(3)
                    p.paragraph_format.line_spacing = 1.1

                    text = cell_data.get_text().strip()
                    r = p.add_run(text)
                    r.font.name = "Montserrat"
                    r.font.size = Pt(9.5)
                    if is_header or cell_data.name == 'th' or cell_data.find(['strong', 'b']):
                        r.font.bold = True

            # Add small spacing after table
            p_space = doc.add_paragraph()
            p_space.paragraph_format.space_before = Pt(2)
            p_space.paragraph_format.space_after = Pt(4)

print("docx_engine.py defined successfully!")
