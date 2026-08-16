import docx
import sys

sys.stdout.reconfigure(encoding='utf-8')
doc = docx.Document(r"C:\Users\Ayush Sood\Downloads\sports final tender.docx")
print("=== PARAGRAPHS 0 to 15 ===")
for i in range(min(16, len(doc.paragraphs))):
    p = doc.paragraphs[i]
    print(f"P[{i}] text: {p.text}")

print("\n=== SECTIONS ===")
for i, s in enumerate(doc.sections):
    print(f"Section {i}: start_type={s.start_type}, header={s.header.paragraphs[0].text if s.header.paragraphs else ''}")
